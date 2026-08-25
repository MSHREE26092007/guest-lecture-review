"""Generate the two sample test documents used by unit tests and the demo.

    python scripts/generate_samples.py

Produces:
    samples/report_good.docx   - a fully compliant guest lecture report
    samples/report_bad.docx    - a report with deliberate violations
"""

import io
import struct
import sys
import zlib
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

SAMPLES_DIR = Path(__file__).resolve().parent.parent / "samples"


def _make_png(width: int = 1, height: int = 1, rgb: tuple = (0, 150, 200)) -> bytes:
    """Build a tiny solid-color PNG without PIL."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        c = tag + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
    row = b"\x00" + bytes(rgb) * width
    idat = zlib.compress(row * height)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _set_normal_style(doc: Document, family: str, size: float, spacing: float) -> None:
    style = doc.styles["Normal"]
    style.font.name = family
    style.font.size = Pt(size)
    style.paragraph_format.line_spacing = spacing


def _set_margins(doc: Document, inches: float) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def _add_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def _add_image(doc: Document, width_in: float = 2.5, alignment=None):
    paragraph = doc.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(_make_png(60, 40)), width=Inches(width_in))
    return paragraph


_SENTENCES = [
    "The resource person opened the session by outlining the motivation behind the topic and its relevance to current industry practice.",
    "Key concepts were introduced step by step, with worked examples shown on the projector to clarify each idea.",
    "Students followed the discussion closely and several asked detailed questions about how the ideas translate to real deployments.",
    "The middle portion of the session demonstrated a live walkthrough of the methodology on a small case study.",
    "Comparisons with alternative approaches were presented, highlighting the trade-offs involved in each design decision.",
    "A short quiz was conducted to check understanding, and the results were discussed openly with the group.",
    "The closing segment summarised the main takeaways and pointed to further reading material shared after the event.",
    "Overall the session was well received and the organisers collected written reflections from the participants.",
]

# Neutral filler that deliberately avoids triggering any required-field or
# policy keyword (no dates, times, venues, institutions, outcomes, etc.).
_BAD_SENTENCES = [
    "This event discussed recent developments in artificial intelligence and its many practical applications.",
    "The session covered machine learning models, data preparation steps, and methods for evaluating results.",
    "Students gained insight into how such systems are built and where they are used in practice.",
    "A live demonstration showed the entire process end to end with a small sample dataset.",
    "The speaker answered questions and gave examples drawn from real world scenarios.",
    "The organisers thanked everyone who attended and shared extra reading material afterwards.",
    "Discussions continued informally after the session had concluded.",
    "Several attendees expressed interest in follow up sessions on related topics.",
]


def _filler(sentences: list[str], n: int) -> list[str]:
    return [sentences[i % len(sentences)] for i in range(n)]


def generate_good_report(path: Path) -> None:
    """Fully compliant report: Calibri 11, 1.15 spacing, 1-inch margins,
    heading styles, page numbers, 3 centered images, styled table, bullets,
    all required fields and policy keywords present."""
    doc = Document()
    _set_normal_style(doc, "Calibri", 11, 1.15)
    _set_margins(doc, 1.0)

    header = doc.sections[0].header
    header.paragraphs[0].text = "Global Technical University"

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = "Page "
    _add_page_number_field(p)

    doc.add_heading("Guest Lecture Report: Artificial Intelligence in Healthcare", level=1)
    _add_image(doc, 2.5, WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("Event Details", level=2)
    for d in [
        "Department of Computer Science",
        "Speaker: Dr. Priya Sharma",
        "Designation: Professor, MedTech University",
        "Organization: MedTech University",
        "Date: 12 March 2026",
        "Time: 10:00 AM to 12:30 PM",
        "Venue: Seminar Hall 2, Main Campus",
        "Faculty Coordinator: Dr. R. Kumar",
    ]:
        doc.add_paragraph(d)

    doc.add_heading("About the Speaker", level=1)
    doc.add_paragraph(_SENTENCES[0])
    doc.add_paragraph(_SENTENCES[1])

    doc.add_heading("Lecture Summary", level=1)
    for s in _filler(_SENTENCES, 32):
        doc.add_paragraph(s)
    _add_image(doc, 2.5, WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("Learning Outcomes", level=1)
    for o in [
        "Understand the core pipeline used in medical image analysis systems.",
        "Explain how diagnostic models are trained and evaluated on clinical data.",
        "Identify ethical and regulatory constraints in healthcare deployments.",
    ]:
        doc.add_paragraph(o, style="List Bullet")

    doc.add_heading("Schedule", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Light Shading Accent 1"
    for row, cells in enumerate([
        ["Time", "Activity", "Resource Person"],
        ["10:00 - 10:15", "Welcome and introduction", "Dr. R. Kumar"],
        ["10:15 - 11:30", "Lecture: AI in Healthcare", "Dr. Priya Sharma"],
        ["11:30 - 12:00", "Live demonstration", "Dr. Priya Sharma"],
        ["12:00 - 12:30", "Q&A and closing", "Dr. Priya Sharma"],
    ]):
        for col, text in enumerate(cells):
            table.rows[row].cells[col].text = text

    doc.add_heading("Student Participation and Attendance", level=1)
    doc.add_paragraph("Attendance: 120 students attended the session.")
    for s in _filler(_SENTENCES, 8):
        doc.add_paragraph(s)

    doc.add_heading("Photos", level=1)
    for s in _filler(_SENTENCES, 3):
        doc.add_paragraph(s)
    _add_image(doc, 2.5, WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_heading("Feedback", level=1)
    doc.add_paragraph("Feedback forms were collected from all participants.")
    doc.add_paragraph("Attendance sheet attached. Budget: INR 25,000 approved.")
    doc.add_paragraph("Invitation letter sent to the speaker. Brochure designed by the event committee.")

    doc.add_heading("Conclusion", level=1)
    for s in _filler(_SENTENCES, 8):
        doc.add_paragraph(s)

    doc.add_heading("Signature", level=1)
    doc.add_paragraph("Signature: Dr. R. Kumar")
    doc.add_paragraph("Faculty Coordinator, Department of Computer Science")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def generate_bad_report(path: Path) -> None:
    """Deliberately non-compliant: Times New Roman 12, 2.0 spacing, 0.5-inch
    margins, no headings, no header/footer/page numbers, one left-aligned
    image, no tables, single page, and most required fields/policy keywords
    missing."""
    doc = Document()
    _set_normal_style(doc, "Times New Roman", 12, 2.0)
    _set_margins(doc, 0.5)

    p = doc.add_paragraph()
    run = p.add_run("Guest Lecture Report: Artificial Intelligence and Society")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("Speaker: John Doe")

    doc.add_page_break()
    _add_image(doc, 1.5, WD_ALIGN_PARAGRAPH.LEFT)

    for s in _filler(_BAD_SENTENCES, 9):
        doc.add_paragraph(s)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def main() -> None:
    good = SAMPLES_DIR / "report_good.docx"
    bad = SAMPLES_DIR / "report_bad.docx"
    generate_good_report(good)
    generate_bad_report(bad)
    print(f"Wrote {good}")
    print(f"Wrote {bad}")


if __name__ == "__main__":
    sys.exit(main())