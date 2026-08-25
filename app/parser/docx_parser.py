"""DOCX parsing using python-docx. Produces the normalized ParsedDocument object."""

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.document import (
    FontRun,
    ImageInfo,
    Margins,
    Paragraph,
    ParsedDocument,
    Section,
    Table,
)

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: None,
}

_HEADING_RE = re.compile(r"^heading\s*(\d+)?$", re.IGNORECASE)
_EXTENT_RE = re.compile(r'<wp:extent cx="([0-9]+)" cy="([0-9]+)"/>')
_EMU_PER_PT = 12700


def _blip_extent_pt(xml: str, axis: str) -> float | None:
    """Return the inline drawing extent (cx or cy) in points, if present."""
    m = _EXTENT_RE.search(xml)
    if not m:
        return None
    emu = int(m.group(1) if axis == "cx" else m.group(2))
    return round(emu / _EMU_PER_PT, 1)


def _docx_page_count(path: Path) -> int:
    """Read docProps/app.xml from the docx zip to get the stored page count."""
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("docProps/app.xml").decode("utf-8", "ignore")
        m = re.search(r"<Pages>(\d+)</Pages>", xml)
        return int(m.group(1)) if m else 0
    except Exception:
        return 0


def _paragraph_details(p, doc) -> Paragraph:
    text = p.text
    style = p.style.name if p.style is not None else None
    alignment = _ALIGN_MAP.get(p.alignment)
    line_spacing = None
    try:
        line_spacing = p.paragraph_format.line_spacing
        if line_spacing is None and p.style is not None:
            line_spacing = p.style.paragraph_format.line_spacing
        if line_spacing is None:
            line_spacing = doc.styles["Normal"].paragraph_format.line_spacing
        if line_spacing is not None and isinstance(line_spacing, float):
            line_spacing = round(line_spacing, 2)
    except Exception:
        pass
    bullet = bool(style and ("list" in style.lower() or "bullet" in style.lower()))
    marker = None
    if bullet:
        marker = "\u2022"  # numPr-based lists render as bullet glyphs
    elif text[:2] in ("- ", "* ", "\u2022 ", "+ "):
        bullet = True
        marker = text[0]
    runs: list[FontRun] = []
    para_style = p.style
    for r in p.runs:
        family = None
        size = None
        try:
            family = r.font.name
            if family is None and para_style is not None:
                family = para_style.font.name
            if family is None:
                family = doc.styles["Normal"].font.name
            if r.font.size is not None:
                size = r.font.size.pt
            elif para_style is not None and para_style.font.size is not None:
                size = para_style.font.size.pt
            elif doc.styles["Normal"].font.size is not None:
                size = doc.styles["Normal"].font.size.pt
        except Exception:
            pass
        runs.append(
            FontRun(
                family=family,
                size_pt=round(size, 1) if size else None,
                bold=bool(r.bold),
                italic=bool(r.italic),
            )
        )
    dominant = None
    if runs:
        sizes = [r.size_pt for r in runs if r.size_pt]
        fams = [r.family for r in runs if r.family]
        dominant = FontRun(
            family=fams[0] if fams else None,
            size_pt=max(sizes) if sizes else None,
            bold=any(r.bold for r in runs),
            italic=any(r.italic for r in runs),
        )
    return Paragraph(
        text=text,
        style=style,
        alignment=alignment,
        line_spacing=line_spacing,
        bullet=bullet,
        list_marker=marker,
        runs=runs,
        dominant_font=dominant,
    )


def _has_page_number_field(doc: Document) -> bool:
    """Look for a PAGE field anywhere in the XML (headers/footers/body)."""
    for part in (doc.part, *[s.header.part for s in doc.sections], *[s.footer.part for s in doc.sections]):
        try:
            xml = part.element.xml
            if re.search(r"instrText[^>]*>\s*PAGE\b", xml, re.IGNORECASE):
                return True
        except Exception:
            continue
    return False


def parse_docx(path: str | Path) -> ParsedDocument:
    path = Path(path)
    doc = Document(str(path))

    sections: list[Section] = []
    paragraphs: list[Paragraph] = []
    tables: list[Table] = []
    images: list[ImageInfo] = []
    image_para_index: list[int] = []
    current: Section | None = None
    current_page = 1

    # Walk the body in document order so paragraphs and tables interleave correctly.
    for child in doc.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            from docx.text.paragraph import Paragraph as DocxParagraph

            p = DocxParagraph(child, doc)
            para = _paragraph_details(p, doc)
            paragraphs.append(para)
            # Images embedded in this paragraph (inline drawings).
            blips = child.xml.count("<a:blip")
            for _ in range(blips):
                images.append(
                    ImageInfo(
                        index=len(images) + 1,
                        page=current_page,
                        width_pt=_blip_extent_pt(child.xml, "cx"),
                        height_pt=_blip_extent_pt(child.xml, "cy"),
                        alignment=para.alignment,
                    )
                )
                image_para_index.append(len(paragraphs) - 1)
            if 'w:type="page"' in child.xml or 'w:type="page" ' in child.xml:
                current_page += 1
            if para.text.strip() == "":
                continue
            if _HEADING_RE.match(para.style or ""):
                m = _HEADING_RE.match(para.style or "")
                level = int(m.group(1) or 1)
                current = Section(heading=para.text, heading_level=level, text=para.text)
                sections.append(current)
            elif current is not None:
                current.text = (current.text + "\n" + para.text).strip()
                current.paragraphs.append(para)
            else:
                sections.append(
                    Section(heading="", heading_level=0, text=para.text, paragraphs=[para])
                )
        elif tag == "tbl":
            from docx.table import Table as DocxTable

            tbl = DocxTable(child, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
            style = tbl.style.name if tbl.style is not None else None
            tables.append(Table(rows=rows, style=style, has_header=len(rows) > 1))
            if current is not None:
                current.text = (current.text + "\n[TABLE: %d rows]" % len(rows)).strip()

    # Margins (first section)
    sec = doc.sections[0]
    margins = Margins(
        top_in=round(sec.top_margin.inches, 2) if sec.top_margin else None,
        bottom_in=round(sec.bottom_margin.inches, 2) if sec.bottom_margin else None,
        left_in=round(sec.left_margin.inches, 2) if sec.left_margin else None,
        right_in=round(sec.right_margin.inches, 2) if sec.right_margin else None,
    )

    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs if p.text.strip())
    footer_text = "\n".join(p.text for p in doc.sections[0].footer.paragraphs if p.text.strip())

    # Captions: the text of the paragraph following each image.
    for img, para_idx in zip(images, image_para_index):
        nxt = para_idx + 1
        while nxt < len(paragraphs) and not paragraphs[nxt].text.strip():
            nxt += 1
        if nxt < len(paragraphs) and paragraphs[nxt].text.strip():
            img.caption = paragraphs[nxt].text.strip()[:120]

    if not images:
        # Anchored (floating) images: scan rels for image parts.
        try:
            for rid in doc.part.rels:
                rel = doc.part.rels[rid]
                if "image" in getattr(rel, "reltype", ""):
                    images.append(
                        ImageInfo(index=len(images) + 1, page=1, width_pt=None, height_pt=None)
                    )
        except Exception:
            pass

    page_count = _docx_page_count(path)
    if not page_count or page_count <= 1:
        # docProps/app.xml is often absent or stale (python-docx ships Pages=1);
        # estimate from text volume (~2000 chars per page at 11pt, 1.15 spacing).
        chars = sum(len(p.text) for p in paragraphs)
        page_count = max(page_count, -(-chars // 2000))
    page_count = max(page_count, 1)

    return ParsedDocument(
        filename=path.name,
        file_type="docx",
        parser="docx",
        ocr_used=False,
        page_count=page_count,
        margins=margins,
        sections=sections,
        tables=tables,
        images=images,
        header_text=header_text,
        footer_text=footer_text,
        page_numbers_present=_has_page_number_field(doc) or bool(re.search(r"\bpage\b", footer_text, re.I)),
        paragraphs=paragraphs,
        text_by_page=[p.text for p in paragraphs],
    )